# batch_md2docx_pycharm.py
# -*- coding: utf-8 -*-

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# =========================
# ✅ 你只需要改这里
# =========================
INPUT_DIR = r"D:\研究生\科研\2.支持型文件\1.AI辅助论文\5.程序汇编\7.文献综述\4.1整体型文献回顾"   # ← 改成你的 md 所在文件夹
RECURSIVE = True        # True=包含子文件夹；False=只扫当前文件夹
OVERWRITE = False       # False=遇到同名docx就跳过；True=覆盖
FONT_NAME = "宋体"      # 可改为 "微软雅黑"
FONT_SIZE = 11          # 正文字号
# =========================


def set_doc_style(doc: Document, font_name=FONT_NAME, font_size=FONT_SIZE):
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_inline_markdown(paragraph, text: str):
    """极简 inline：**bold**、*italic*、`code`"""
    pattern = r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_md_table(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if not line.strip().startswith("|"):
            break
        # 跳过对齐分隔线（---）
        if re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line.strip()):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_text: str, out_path: Path):
    doc = Document()
    set_doc_style(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # 代码块 ``` 开关
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line.rstrip("\n"))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # 表格（|...|）
        if line.strip().startswith("|"):
            rows, next_i = parse_md_table(lines, i)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, r in enumerate(rows):
                    for c_idx in range(cols):
                        table.cell(r_idx, c_idx).text = r[c_idx] if c_idx < len(r) else ""
                i = next_i
                continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h_level = level if level <= 4 else 4
            doc.add_heading(text, level=h_level)
            i += 1
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, m.group(2).strip())
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, m.group(2).strip())
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline_markdown(p, line.strip())
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def batch_convert(folder: Path):
    pattern = "**/*.md" if RECURSIVE else "*.md"
    md_files = sorted(folder.glob(pattern))

    if not md_files:
        print(f"⚠️ 没找到 .md 文件：{folder}")
        return

    ok = skipped = fail = 0

    for md_path in md_files:
        docx_path = md_path.with_suffix(".docx")

        if docx_path.exists() and not OVERWRITE:
            print(f"⏭️ 跳过(已存在): {docx_path.name}")
            skipped += 1
            continue

        try:
            md_text = md_path.read_text(encoding="utf-8")
            md_to_docx(md_text, docx_path)
            print(f"✅ 转换: {md_path.name} -> {docx_path.name}")
            ok += 1
        except Exception as e:
            print(f"❌ 失败: {md_path} | {e}")
            fail += 1

    print(f"\n🎉 完成：成功 {ok}，跳过 {skipped}，失败 {fail}")
    print(f"📁 输出位置：仍在原目录，同名 .docx 文件已生成")


if __name__ == "__main__":
    folder = Path(INPUT_DIR).expanduser().resolve()
    if not folder.exists() or not folder.is_dir():
        print(f"❌ 输入文件夹不存在：{folder}")
    else:
        batch_convert(folder)